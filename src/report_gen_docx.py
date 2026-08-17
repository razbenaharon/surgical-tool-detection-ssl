"""Generate the HW1 DOCX report from the same content as report_gen.py.

Usage:
    python src/report_gen_docx.py --root . --out reports/HW1_report_final.docx
"""
import argparse
import json
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


def add_figure(doc, path, caption, width_inches=5.5):
    p = Path(path)
    if not p.exists():
        doc.add_paragraph(f"[missing figure: {path}]")
        return
    doc.add_picture(str(p), width=Inches(width_inches))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_para = doc.add_paragraph(caption)
    cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap_para.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(128, 128, 128)
        run.font.italic = True


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(7.5)
    # Data
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(7.5)
    return table


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--root", default=".")
    ap.add_argument("--metrics", default="reports/metrics.json")
    ap.add_argument("--out", default="reports/HW1_report_final.docx")
    args = ap.parse_args()
    root = Path(args.root)

    M = json.loads(Path(args.metrics).read_text()) if Path(args.metrics).exists() else {}
    eda = {}
    eda_path = root / "reports/eda/dataset_summary.json"
    if eda_path.exists():
        eda = json.loads(eda_path.read_text())

    doc = Document()

    # Set narrow margins
    for section in doc.sections:
        section.top_margin = Cm(1.3)
        section.bottom_margin = Cm(1.3)
        section.left_margin = Cm(1.3)
        section.right_margin = Cm(1.3)

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(9)

    tr = eda.get("train", {})
    va = eda.get("val", {})

    # Title
    title = doc.add_heading('Surgical Tool Detection with Semi-Supervised Learning', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph('Computer Vision \u2014 Surgical Applications, HW1')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    authors = doc.add_paragraph('Authors: Raz Ben-Aharon and Shalev Manassen (shalevmanassen@gmail.com)')
    authors.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ===== 1. EDA =====
    doc.add_heading('1. Exploratory Data Analysis', level=1)
    doc.add_paragraph(
        f'The labeled set contains only {tr.get("n_images", 61)} training and '
        f'{va.get("n_images", 10)} validation images (<100 total), all native 4K '
        f'(3840\u00d72160). Each frame comes from a leg-suturing surgery video; every box '
        f'encloses a gloved hand together with whatever it holds, and the class label is '
        f'defined by what the hand holds: Empty (no tool), Tweezers, or Needle_driver.')

    doc.add_heading('1.1 Visualization of Some Images', level=2)
    doc.add_paragraph(
        'Fig. 1 shows a random sample of labeled training frames with their '
        'ground-truth boxes drawn (green = Empty, blue = Tweezers, red = Needle_driver).')
    add_figure(doc, root / "reports/eda/samples_grid.png",
               "Fig 1. Nine labeled frames with ground-truth boxes overlaid.", 5.0)

    doc.add_heading('1.2 Insights from simply \u201clooking\u201d at the data', level=2)
    doc.add_paragraph(
        'Looking through the labeled frames and the raw videos suggests: (i) a static '
        'top-down camera view with hands in a consistent region; (ii) large boxes (15\u201335% '
        'of frame width, covering hand-plus-tool); (iii) a visually difficult scene with '
        'specular highlights, blood, and similar-coloured drapes; (iv) 1\u20132 boxes per frame '
        'with no background-only images; (v) the OOD video has different lighting/camera.')

    doc.add_heading('1.3 Data distribution analysis', level=2)
    add_table(doc,
              ["Split", "Images", "Boxes", "Empty", "Tweezers", "Needle_driver", "Boxes/img"],
              [["train", tr.get("n_images", 61), tr.get("n_boxes", 135),
                tr.get("class_counts", {}).get("Empty", 26),
                tr.get("class_counts", {}).get("Tweezers", 54),
                tr.get("class_counts", {}).get("Needle_driver", 55),
                f'{tr.get("boxes_per_image_mean", 2.21):.2f}'],
               ["val", va.get("n_images", 10), va.get("n_boxes", 22),
                va.get("class_counts", {}).get("Empty", 2),
                va.get("class_counts", {}).get("Tweezers", 10),
                va.get("class_counts", {}).get("Needle_driver", 10),
                f'{va.get("boxes_per_image_mean", 2.2):.2f}']])
    doc.add_paragraph(
        'Empty is the clear minority class (26/135 \u2248 19% of train boxes), while '
        'Tweezers and Needle_driver are roughly balanced.')
    add_figure(doc, root / "reports/eda/class_distribution.png",
               "Fig 2. Class distribution by split.", 3.0)
    add_figure(doc, root / "reports/eda/box_size_distribution.png",
               "Fig 3. Box size distribution and width vs. height.", 5.0)

    # ===== 2. Experiments =====
    doc.add_page_break()
    doc.add_heading('2. Experiments', level=1)
    doc.add_paragraph(
        'Overview. Following the assignment\u2019s SSL guideline, we: (1) fine-tuned a '
        'COCO-pretrained YOLO11s detector on the 61 labeled training images; '
        '(2) pseudo-labeled the ID and OOD videos with various filtering policies; '
        '(3) retrained and evaluated each variant. Additionally, we ran a supervised '
        'hyperparameter sweep over resolution, learning rate and model scale. '
        'The result was unexpected: the original supervised Baseline produced the best '
        'OOD behavior, while pseudo-label self-training degraded OOD performance '
        'through confirmation bias (\u00a73).')

    doc.add_heading('2.1 Data loading, pre-processing and cleaning', level=2)
    doc.add_paragraph(
        'Loading. Images and YOLO-format labels are read through Ultralytics\u2019 standard '
        'dataloader. For SSL rounds, build_ssl_dataset.py combines labeled and pseudo-labeled '
        'images while keeping the 10-image validation set unchanged. '
        'Pre-processing. Images are letterbox-resized to the target resolution (640\u00d7640 default). '
        'Split. The original train/val split (61/10 images) was kept as-is.')
    doc.add_paragraph(
        'Validation annotation audit. Manual inspection identified two annotation artifacts in '
        'image ff8c22da-output_0182.png: a 14\u00d714-pixel Needle_driver box on dark background '
        '(row 2) and a 10\u00d719-pixel Tweezers box at a glove/background boundary (row 4). '
        'With only 10 validation images and \u224822 boxes, these represent \u22489% of the evaluation '
        'set. Two validation snapshots were created: val_original (byte-identical to official '
        'data) and val_clean (only those two removed). Baseline: original mAP50-95 = 0.754, '
        'clean mAP50-95 = 0.809. Epoch selection replay still selected epoch 128.')

    doc.add_heading('2.2 Training techniques', level=2)
    doc.add_paragraph(
        'Transfer learning. All experiments fine-tune YOLO11s (9.4M params) from COCO weights. '
        'Optimizer & schedule. optimizer=auto resolved to AdamW (momentum 0.9, weight decay 5e-4) '
        'with effective initial LR \u22480.001429 on a cosine-annealed schedule. Note: the reported '
        'lr0=0.01 was overridden by the auto-optimizer. Batch 16, imgsz 640, mixed-precision. '
        'Data augmentation (moderate): mosaic (p=1.0), HSV jitter, horizontal flip (p=0.5), '
        'mild translate (0.1) and scale (0.5). No aggressive augmentation \u2014 it caused divergence.')
    doc.add_paragraph(
        'Semi-supervised pseudo-labeling. Original SSL: ID videos yielded 702/720 frames, '
        '1,690 boxes at mean conf 0.876. OOD videos yielded 667/713 frames, 2,147 boxes '
        'at mean conf 0.732. 71.7% of the OOD pseudo-label boxes were classified as '
        'Needle_driver, revealing a strong class skew. Manual inspection additionally showed '
        'recurring false-positive Needle_driver detections on blood-stained gauze and '
        'background regions, suggesting that systematic teacher errors were being reinforced '
        'during self-training. Combined with 61 real images, the OOD training set had '
        '1,369 pseudo-labeled images (22:1 pseudo:real ratio).')

    doc.add_heading('2.3 Regularization', level=2)
    doc.add_paragraph(
        'Weight decay 5e-4; early stopping (patience 50); moderate data augmentation; '
        'COCO pretrained backbone as prior. Dropout is not used (dropout=0.0).')

    doc.add_heading('2.4 Hyperparameter tuning', level=2)
    doc.add_paragraph(
        'Systematic supervised sweep: resolution (640/960/1280), model scale (YOLO11s/m), '
        'LR (0.001/0.003/0.01), augmentation (moderate/weak). Key findings: '
        'LR 0.003/0.01 caused collapse/NaN; only 0.001 was stable. YOLO11m did not improve '
        'over YOLO11s. Supervised 960 achieved the highest clean mAP50-95 = 0.854 but '
        'showed recurring gauze false positives on OOD video. Higher ID validation mAP '
        'did not guarantee better OOD generalization.')
    doc.add_paragraph(
        'Pseudo-label policies tested: confidence-only (P1\u2013P4), sanity constraints, '
        'per-frame caps, class-specific thresholds, temporal persistence. P4 retained only '
        '69 frames at mean conf 0.938 \u2014 model still produced substantial false positives, '
        'showing confidence \u2260 correctness. Strict OOD thresholds (box 0.85/frame 0.92) '
        'accepted zero OOD frames. Relaxed combined approach retained 123 OOD frames '
        'but made the model too conservative (21% empty frames).')

    doc.add_heading('2.5 Train + valid loss graph', level=2)
    add_figure(doc, root / "reports/curves/loss_curves.png",
               "Fig 4. Training and validation loss vs. epoch for the Baseline and SSL rounds.", 5.5)
    doc.add_paragraph(
        'Baseline shows early instability (epochs 5\u201331) during LR warmup, then recovers. '
        'SSL rounds (warm-started) do not show this instability.')

    doc.add_heading('2.6 Train + valid mAP graphs', level=2)
    add_figure(doc, root / "reports/curves/map_curves.png",
               "Fig 5. Validation mAP@50 and mAP@50-95 vs. epoch (official val set).", 5.5)

    add_table(doc,
              ["Model", "Epoch", "Imgs", "mAP50-95\norig", "mAP50-95\nclean", "OOD behavior"],
              [["Baseline (final)", "128", "61", "0.754", "0.809", "Best balance"],
               ["ssl_id", "40", "763", "0.782", "0.837", "87% >2 det"],
               ["ssl_ood", "4", "1430", "0.740", "0.792", "90% >2 det"],
               ["Sup. 960", "103", "61", "0.797", "0.854", "Gauze FP"],
               ["Sup. 1280", "76", "61", "0.751", "0.805", "Flicker"],
               ["ID P4", "11", "130", "0.749", "0.800", "81% >2 det"],
               ["ID temporal", "18", "627", "0.767", "0.820", "16% empty"],
               ["ID+OOD comb.", "24", "750", "0.721", "0.775", "21% empty"]])
    cap = doc.add_paragraph(
        'Table 1. Key experiment candidates with official and cleaned validation mAP50-95 '
        'and OOD video behavior.')
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap.runs:
        run.font.size = Pt(8)
        run.font.italic = True

    # ===== 3. Discussion and Conclusions =====
    doc.add_page_break()
    doc.add_heading('3. Discussion and Conclusions', level=1)
    doc.add_paragraph(
        'Since the OOD video has no ground-truth labels, OOD performance is assessed '
        'using automated diagnostics and manual visual inspection.')

    doc.add_heading('Original SSL failure', level=2)
    doc.add_paragraph(
        'Naive pseudo-label self-training degraded OOD behavior in our experiments. '
        'The OOD pseudo-label boxes were 71.7% classified as Needle_driver, revealing '
        'a strong class skew. Furthermore, 1,369 pseudo-labeled images overwhelmed 61 real '
        'images. The ssl_ood model (epoch 4, mAP50-95 = 0.740) produced 89.6% of long-video '
        'frames with >2 detections (Baseline: 6.6%). The Baseline outperformed the old '
        'ssl_ood model on cleaned validation mAP and, more importantly, showed substantially '
        'better OOD video behavior.')

    doc.add_heading('Pseudo-label policy findings', level=2)
    doc.add_paragraph(
        '(i) Confidence alone is insufficient: P4 (69 frames, mean conf 0.938) still '
        'produced 81.3% >2 frames. (ii) Temporal consistency helps: P3+sanity+temporal '
        '(0.5% >2 frames, 18 class switches) but at recall cost (16% empty frames). '
        '(iii) Strict OOD thresholds rejected everything; relaxed approach (123 frames) '
        'made the model too conservative.')

    doc.add_heading('Finalist comparison and final model selection', level=2)
    add_table(doc,
              ["Finalist", "det/fr", ">2 %", "empty %", "1-frame tracks", "switches"],
              [["Baseline / 0.60", "1.70", "6.6", "5.9", "145", "69"],
               ["ID temporal / 0.40", "1.38", "0.5", "16.0", "92", "18"],
               ["Sup. 1280 / 0.50", "1.49", "5.1", "10.5", "288", "82"]])
    cap2 = doc.add_paragraph(
        'Table 2. Diagnostics on full 180s OOD video (5,395 frames) at each finalist\u2019s '
        'inference threshold.')
    cap2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap2.runs:
        run.font.size = Pt(8)
        run.font.italic = True
    doc.add_paragraph(
        'Manual review of identical OOD segments focused on false positives, missed '
        'detections, class stability, and temporal stability. Baseline (conf=0.60): best '
        'overall visual behavior \u2014 reliable hand detection without recurring gauze false '
        'positives. ID temporal (conf=0.40): very stable but sometimes missed real hands. '
        'Sup. 1280 (conf=0.50): reasonable coverage but excessive temporal flicker (288 '
        'one-frame tracks). The supervised Baseline at inference confidence 0.60 was '
        'selected as the final model.')

    doc.add_heading('Key lessons', level=2)
    lessons = [
        '(1) Pseudo-label quality > quantity \u2014 systematic errors were amplified.',
        '(2) Self-training can create confirmation bias \u2014 gauze \u2192 Needle_driver reinforced.',
        '(3) Confidence \u2260 correctness \u2014 P4 at mean conf 0.938 still failed.',
        '(4) Temporal consistency useful but insufficient \u2014 reduces FP but also recall.',
        '(5) ID validation mAP \u2260 OOD quality \u2014 Sup. 960 (mAP 0.854) had gauze FP.',
        '(6) Small val sets sensitive to noise \u2014 2 boxes changed mAP by \u22487%.',
        '(7) Simplest model can win \u2014 complexity must be validated, not assumed.',
    ]
    for lesson in lessons:
        doc.add_paragraph(lesson)

    doc.add_heading('Conclusion', level=2)
    doc.add_paragraph(
        'The supervised Baseline was selected as the final detector because it produced '
        'the best overall OOD video behavior, despite not achieving the highest validation '
        'mAP. Naive pseudo-label self-training introduced confirmation bias and severe false '
        'positives. Stricter confidence thresholds alone were insufficient, while sanity '
        'and temporal filtering improved prediction stability but reduced recall. The '
        'experiments therefore show that higher validation metrics and larger pseudo-labeled '
        'datasets did not necessarily translate into better OOD generalization.')

    out_path = root / args.out
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(f"wrote {args.out}")


if __name__ == "__main__":
    main()
